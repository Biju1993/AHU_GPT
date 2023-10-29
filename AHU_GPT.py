import streamlit as st
import requests
import pandas as pd
import pythermalcomfort
from pythermalcomfort.models import pmv_ppd
from pythermalcomfort.utilities import v_relative, clo_dynamic
from fpdf import FPDF
import requests
import psychrolib
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO



def get_location():
    try:
        # Make a request to ipinfo.io to get the location information based on the user's IP address
        response = requests.get('http://ipinfo.io')
        data = response.json()

        # Parse the location data
        ip = data.get('ip')
        city = data.get('city')
        region = data.get('region')
        country = data.get('country')
        location = f"IP: {ip}, City: {city}, Region: {region}, Country: {country}"
        # return location
        return city

    except Exception as e:
        return f"Error: {e}"

def home():
    st.title("Welcome to AI trained report generator....")
    st.write("Choose a page from the sidebar to get started....")
    st.write("The machine learning model is for educational understanding only...")
    st.write("The features incorporated are listed below")
    # List of items for bulleted points
    bullet_points = ["Auto data selection based on usage",
                     "Site location selection based on ip address",
                     "User input of data and auto select certain parameter", 
                     "Interactive checklist parameters inputs",
                     "Automatic detection of real time weather data using API",
                     "Automatic Air velocity input prompt based on number of input",
                     "Automatic Psychrometric calculation based on user input",
                     "Automatic Plotting of Psychrometric data at location based sea level",
                     "Automatic TR and IKW calculation",
                     "ASHRAE 55 PPD calculation",
                     "ASHRAE 55 PMV Calculation",
                     "Percentage deviation from selected and measured onsite reading",
                     "Downloading reports in word and pdf format"]

    # Display bulleted points using st.markdown()
    bullet_points_html = "<ul>" + "".join([f"<li>{item}</li>" for item in bullet_points]) + "</ul>"
    st.markdown(bullet_points_html, unsafe_allow_html=True)

def form():
    st.title("AHU PRE COMMISSIONING REPORT")
    st.header("AHU SELECTION DATA")
    DATE_OF_INSPECTION=str(st.date_input("SELECT DATE OF TESTING"))
    AHU_TAG_NAME=st.text_input("ENTER NAME TAG/AHU NAME/IDENTITY NUMBER")
    AHU_DESIGN_CFM=st.number_input("ENTER THE AHU DESIGN CFM IN WHOLE NUMBER")
    AHU_DESIGN_ESP_IN_MMWG=st.number_input("ENTER THE AHU EXTERNAL STATIC PRESSURE IN MM of WG")
    SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0, 25.0)
    SELECTED_INDOOR_RELATIVE_HUMIDITY=st.slider("SELECTED_INDOOR_RELATIVE_HUMIDITY", 10.0, 100.0, 25.0,step=0.5)
    SELECTED_ZONE_RELATIVE_AIR_VELOCITY=st.slider("SELECTED_ZONE_RELATIVE_AIR_VELOCITY", 0.1, 1.0,0.25,step=0.05)
    METABOLIC_EQUIVALENT_MER=st.slider("METABOLIC_EQUIVALENT_MER", 0.7, 4.0,1.0,step=0.1)
    CLOTHING_INSULATION_CLO=st.slider("METABOLIC_EQUIVALENT_MER", 0.36, 1.0,0.61,step=0.05)
    CURRENT_LOCATION=get_location()
    SITE_LOCATION_OF_THE_UNIT=st.selectbox(" SITE_LOCATION_OF_THE_UNIT",[CURRENT_LOCATION,'bangalore','mumbai', 'chennai'],)
    
    # return AHU_TAG_NAME,AHU_DESIGN_CFM
    #Map feature:
    
    weather_api_endpoint = "http://api.openweathermap.org/data/2.5/weather"
    api_key = "8af1a752ec6e57631f899729ded7eba5"  # Replace this with your actual OpenWeatherMap API key

    params = {'q':  SITE_LOCATION_OF_THE_UNIT,
            'appid': api_key,
            'units': 'metric'  }
    response = requests.get(weather_api_endpoint, params=params)
    if response.status_code == 200:
        weather_data = response.json()
        main = weather_data['main']
        weather = weather_data['weather'][0]['description']
        temperature = main['temp']
        atm= main['pressure']
        humidity = main['humidity']
        lon=(weather_data['coord']['lon'])
        lat=(weather_data['coord']['lat'])
        win=(weather_data['wind']['speed'])

    else:
        print("Error fetching weather data.")
    
    # st.map(latitude=12.976, longitude=77.6033)
    # st.map(latitude="12.976", longitude="77.603")
    
    f11q={SITE_LOCATION_OF_THE_UNIT: 0}
    q = pd.DataFrame(
        [[lat,lon]],
        columns=['lat', 'lon'])
    st.map(q.loc[[f11q.get(SITE_LOCATION_OF_THE_UNIT)]])
    
    AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW=st.number_input("AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW")
    SELECTED_MOTOR_EFFICIENCY=st.slider("SELECTED_MOTOR_EFFICIENCY", 50.0, 100.0,90.0,step=0.5)
    SELECTED_MOTOR_POWER_FACTOR=st.slider("SELECTED_MOTOR_POWER_FACTOR", 0.7, 1.0,0.89,step=0.05)
    SELECTED_RATED_CURRENT=st.number_input("SELECTED_RATED_CURRENT")
    SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS",value=16)
    SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIU",value=7)
    SELECTED_CHILLED_WATER_FLOWRATE_USGPM=st.number_input("SELECTED_CHILLED_WATER_FLOWRATE_USGPM",value=50)
    
    
    
    st.subheader("GENERAL OBSERVATION")
    
    Permanent_labels_affixed_for_units=st.radio("Permanent_labels_affixed_for_units",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Casing_condition_good_no_dents_leaks_door_gaskets_installed=st.radio("Casing_condition_good_no_dents_leaks_door_gaskets_installed",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Access_doors_close_tightly__no_leaks=st.radio("Access_doors_close_tightly__no_leaks",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Vibration_isolation_installed_and_released_from_shipping_locks=st.radio("Vibration_isolation_installed_and_released_from_shipping_locks",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Maintenance_access_acceptable_for_unit_and_components=st.radio("Maintenance_access_acceptable_for_unit_and_components",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Thermal_insulation_properly_done_according_to_specification=st.radio("Thermal_insulation_properly_done_according_to_specification",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Instrumentation_installed_according_to_specification=st.radio("Instrumentation_installed_according_to_specification",["YES","NO","NOT APPLICABLE"],horizontal=True)
    
    
    st.subheader("FANS & DAMEPERS")
    
    Fan_is_free_to_rotate_and_in_correct_direction_as_per_approved_shop_drawings=st.radio("Fan_is_free_to_rotate_and_in_correct_direction_as_per_approved_shop_drawings",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Supply_fan_and_motor_alignment_corrects=st.radio("Supply_fan_and_motor_alignment_corrects",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Supply_fan_area_cleans=st.radio("Supply_fan_area_cleans",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Supply_fan_and_motor_properly_lubricateds=st.radio("Supply_fan_and_motor_properly_lubricateds",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Filter_pressure_differential_measuring_device_installed_and_functionals=st.radio("Filter_pressure_differential_measuring_device_installed_and_functionals",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Smoke_and_fire_dampers_installed_properlys=st.radio("Smoke_and_fire_dampers_installed_properly",["YES","NO","NOT APPLICABLE"],horizontal=True)
    
    
    st.subheader("ELECTRICAL & CONTROLS")    
    
    Microprocessor_based_Control_panel_with_Display=st.radio("Microprocessor_based_Control_panel_with_Display",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Emergency_stop_button=st.radio("Emergency_stop_button",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Inspection_Lights=st.radio("Inspection_Lights",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Limit_switch=st.radio("Limit_switch",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Air_temp_in_andOut_sensors=st.radio("Air_temp_in_andOut_sensors",["YES","NO","NOT APPLICABLE"],horizontal=True)
    CHWS_and_R_sensors=st.radio("CHWS_and_R_sensors",["YES","NO","NOT APPLICABLE"],horizontal=True)
    DPT_switch_across_fan_section=st.radio("DPT_switch_across_fan_section",["YES","NO","NOT APPLICABLE"],horizontal=True)
    DP_switch_across_filter=st.radio("DP_switch_across_filter",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Co2_sensor=st.radio("Co2_sensor",["YES","NO","NOT APPLICABLE"],horizontal=True)
    Proper_grounding_installed_for_components_and_unit=st.radio("Proper_grounding_installed_for_components_and_unit",["YES","NO","NOT APPLICABLE"],horizontal=True)
    
    # PRE CX CHECKLIST DATA

    pc=[
    "Permanent labels affixed for units",
    "Casing condition good: no dents, leaks, door gaskets installed",
    "Access doors close tightly - no leaks",
    "Vibration isolation installed & released from shipping locks",
    "Maintenance access acceptable for unit and components",
    "Thermal insulation properly done according to specification.",
    "Instrumentation installed according to specification"]

    pcind=[Permanent_labels_affixed_for_units,
    Casing_condition_good_no_dents_leaks_door_gaskets_installed,
    Access_doors_close_tightly__no_leaks,
    Vibration_isolation_installed_and_released_from_shipping_locks,
    Maintenance_access_acceptable_for_unit_and_components,
    Thermal_insulation_properly_done_according_to_specification,
    Instrumentation_installed_according_to_specification,]


    pc2=[
    "Fan is free to rotate and in correct direction as per approved shop drawing",
    "Supply fan and motor alignment correct",
    "Supply fan area clean",
    "Supply fan and motor properly lubricated",
    "Filter pressure differential measuring device installed and functional",
    "Smoke and fire dampers installed properly (proper location, access doors, appropriate ratings verified)"]

    pc21=[Fan_is_free_to_rotate_and_in_correct_direction_as_per_approved_shop_drawings,
    Supply_fan_and_motor_alignment_corrects,
    Supply_fan_area_cleans,
    Supply_fan_and_motor_properly_lubricateds,
    Filter_pressure_differential_measuring_device_installed_and_functionals,
    Smoke_and_fire_dampers_installed_properlys]




    ## "ELECTRICAL & CONTROLS"
    pc3=["Microprocessor based Control panel with Display",
    "Emergency stop button",
    "Inspection Lights",
    "Limit switch",
    "Air temp in/Out sensors",
    "CHWS/R sensors",
    "DPT switch across fan section",
    "DP switch across filter",
    "Co2 sensor",
    "Proper grounding installed for components and unit"]
    
    st.subheader("FILTERS")    

    pc31=[Microprocessor_based_Control_panel_with_Display,
    Emergency_stop_button,
    Inspection_Lights,
    Limit_switch,
    Air_temp_in_andOut_sensors,
    CHWS_and_R_sensors,
    DPT_switch_across_fan_section,
    DP_switch_across_filter,
    Co2_sensor,
    Proper_grounding_installed_for_components_and_unit
    ]

    COMMISSIONING_FILTER_FOR_PROVIDED=st.radio("COMMISSIONING_FILTER_FOR_PROVIDED",["YES","NO","NOT APPLICABLE"],horizontal=True)
    PRE_FILTER_INSTALLED_MERV_8=st.radio("PRE_FILTER_INSTALLED_MERV_8",["YES","NO","NOT APPLICABLE"],horizontal=True)
    FINE_FILTER_INSTALLED_MERV_14=st.radio("FINE_FILTER_INSTALLED_MERV_14",["YES","NO","NOT APPLICABLE"],horizontal=True)
    UVGI_INSTALLED_AS_PER_SPECIFICATION=st.radio("UVGI_INSTALLED_AS_PER_SPECIFICATION",["YES","NO","NOT APPLICABLE"],horizontal=True)

    pc4=[
    "COMMISSIONING_FILTER_FOR_PROVIDED",
    "PRE_FILTER_INSTALLED_MERV-8",
    "FINE_FILTER_INSTALLED_MERV-14",
    "UVGI_INSTALLED_AS_PER_SPECIFICATION"
    ]
    pc41=[
    COMMISSIONING_FILTER_FOR_PROVIDED,
    PRE_FILTER_INSTALLED_MERV_8,
    FINE_FILTER_INSTALLED_MERV_14,
    UVGI_INSTALLED_AS_PER_SPECIFICATION
    ]
    
    class PDFWithTable(FPDF):
        def header(self):
            # Add a header if needed
            pass

        def footer(self):
            # Add a footer if needed
            pass

        def create_table(self, header, data):
            # Set column widths based on the number of columns
            col_widths = [self.w / len(header)] * len(header)

            # Set font for table header
            self.set_font("Arial", "B", 8)

            # Add table header
            for i in range(len(header)):
                self.cell(col_widths[i]-5, 10, header[i], 1)
            self.ln()

            # Set font for table rows
            self.set_font("Arial", size=8)

            # Add table rows
            for row in data:
                for i in range(len(row)):
                    self.cell(col_widths[i]-5, 10, str(row[i]), 1)
                self.ln()

    ATTENDEES_01 =st.selectbox("SELECT ATTENDEES_01",["CLIENT", "T&C CONSULTANT", "PMC","OEM","FM TEAM"] )
    ATTENDEES_02 =st.selectbox("SELECT ATTENDEES_02",["CLIENT", "T&C CONSULTANT", "PMC","OEM","FM TEAM"] )
    ATTENDEES_03 =st.selectbox("SELECT ATTENDEES_03",["CLIENT", "T&C CONSULTANT", "PMC","OEM","FM TEAM"] )

    ATTENDEES_01_NAME = st.text_input("ENTER ATTENDEES_01_NAME")
    ATTENDEES_02_NAME = st.text_input("ENTER ATTENDEES_02_NAME")
    ATTENDEES_03_NAME = st.text_input("ENTER ATTENDEES_03_NAME")

# Example usage
# if __name__ == "__main__":
    # Sample data for the table
    table_header = [ATTENDEES_01, ATTENDEES_02, ATTENDEES_03]
    table_data = [
        [ATTENDEES_01_NAME , ATTENDEES_02_NAME , ATTENDEES_03_NAME ],
        ["SIGNATURE","SIGNATURE","SIGNATURE"]
    ]
    pdf = PDFWithTable()
    pdf.add_page()

    pdf.set_font("Times",'B', size=12)


    pdf.cell(200, 8, txt="AHU PRE COMMISSIONING REPORT", ln=True, align='C')
    pdf.ln(6)
    pdf.set_fill_color(200, 220, 255)
    pdf.cell(175, 6, txt="DATE OF INSPECTION"+" =  "  +DATE_OF_INSPECTION, ln=True, align='L', border=1)
    pdf.ln(4)
    pdf.cell(175, 6, txt="AHU TAG NAME"+" =  "  +AHU_TAG_NAME, ln=True, align='L', border=1)
    pdf.ln(4)
    pdf.cell(175, 6, txt="AHU DESIGN CFM"+" =  "  +str(AHU_DESIGN_CFM), ln=True, align='L', border=1)
    pdf.ln(4)
    pdf.cell(175, 6, txt="AHU DESIGN ESP IN MMWG"+" =  "  +str(AHU_DESIGN_ESP_IN_MMWG), ln=True, align='L', border=1)
    pdf.ln(4)




    pcx=[SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS,
    SELECTED_INDOOR_RELATIVE_HUMIDITY,
    SELECTED_ZONE_RELATIVE_AIR_VELOCITY,
    METABOLIC_EQUIVALENT_MER,
    CLOTHING_INSULATION_CLO,
    SITE_LOCATION_OF_THE_UNIT,
    SELECTED_MOTOR_EFFICIENCY,
    SELECTED_MOTOR_POWER_FACTOR,
    SELECTED_RATED_CURRENT,
    SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS,
    SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS,
    SELECTED_CHILLED_WATER_FLOWRATE_USGPM]

    pcxi=["SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_INDOOR_RELATIVE_HUMIDITY",
    "SELECTED_ZONE_RELATIVE_AIR_VELOCITY",
    "METABOLIC_EQUIVALENT_MER",
    "CLOTHING_INSULATION_CLO",
    "SITE_LOCATION_OF_THE_UNIT",
    "SELECTED_MOTOR_EFFICIENCY",
    "SELECTED_MOTOR_POWER_FACTOR",
    "SELECTED_RATED_CURRENT",
    "SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_CHILLED_WATER_FLOWRATE_USGPM"]

    for i in range(len(pcx)-1):
        pdf.set_font("Times",'B', size=12)
        pdf.cell(175, 6, txt=pcxi[i]+" =  "  +str(pcx[i]), ln=True, align='L', border=1)
        pdf.ln(4)

    pdf.set_font("Times",'B', size=10)
    pdf.cell(180, 6, txt="GENERAL OBSREVATION", ln=True, align='C', border=1)
    pdf.ln(5)

    for i in range(len(pc)):
        pdf.set_font("Times", size=8)
        pdf.cell(180, 5, txt=pc[i].upper()+" =  "  +pcind[i], ln=True, align='L', border=1)
        pdf.ln(2)

    pdf.set_font("Times",'B', size=10)
    pdf.cell(180, 6, txt="FANS & DAMPERS OBSERVATION", ln=True, align='C', border=1)
    pdf.ln(5)

    for i in range(len(pc2)-1):
        pdf.set_font("Times", size=8)
        pdf.cell(180, 5, txt=pc2[i].upper()+" =  "  +pc21[i], ln=True, align='L', border=1)
        pdf.ln(2)

    pdf.set_font("Times",'B', size=10)
    pdf.cell(180, 6, txt="ELECTRICAL & CONTROLS", ln=True, align='C', border=1)
    pdf.ln(5)

    for i in range(len(pc3)-1):
        pdf.set_font("Times", size=8)
        pdf.cell(180, 5, txt=pc3[i].upper()+" =  "  +str(pc31[i]), ln=True, align='L', border=1)
        pdf.ln(2)

    pdf.set_font("Times",'B', size=10)
    pdf.cell(180, 6, txt="FILTERS", ln=True, align='C', border=1)
    pdf.ln(5)

    for i in range(len(pc4)-1):
        pdf.set_font("Times", size=8)
        pdf.cell(180, 5, txt=pc4[i].upper()+" =  "  +str(pc41[i]), ln=True, align='L', border=1)
        pdf.ln(2)

    pdf.create_table(table_header, table_data)

    # Save the pdf with name .pdf


# The above code is a Python code snippet that is part of a larger program. It is using the Streamlit
# library to create a button that, when clicked, generates a pre-commissioning report in PDF format.
    if st.button("Generate Pre Commissioning Report"):
        # Save the pdf with name .pdf
        pdf_file_path = "AHU_PRE_CX_REPORT.pdf"
        pdf.output(pdf_file_path)
        # Provide download link for generated PDF
        st.success("Report generated successfully!")
        st.download_button(
            label="Download Report",
            data=open(pdf_file_path, "rb").read(),
            file_name=AHU_TAG_NAME+"PRE_CX_REPORT.pdf",
        )

def data_display():
    st.title("Data Display Page")
    # Display data or charts here
    st.write("UNDER CONSTRUCTION.")
    picture = st.camera_input("Take a picture of the snag using mobile and GPS location on")
    st.image("https://img.freepik.com/premium-vector/construction-sign-label_24886-506.jpg")

    if picture:
        st.image(picture)

def loop():
    st.title("AHU PERFORMANCE REPORT")
    st.header("AHU SELECTION DATA")
    DATE_OF_INSPECTION=str(st.date_input("SELECT DATE OF TESTING"))
    AHU_TAG_NAME=st.text_input("ENTER NAME TAG/AHU NAME/IDENTITY NUMBER")
    AHU_DESIGN_CFM=st.number_input("ENTER THE AHU DESIGN CFM IN WHOLE NUMBER")
    AHU_DESIGN_ESP_IN_MMWG=st.number_input("ENTER THE AHU EXTERNAL STATIC PRESSURE IN MM of WG")
    SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0, 25.0)
    SELECTED_INDOOR_RELATIVE_HUMIDITY=st.slider("SELECTED_INDOOR_RELATIVE_HUMIDITY", 10.0, 100.0, 25.0,step=0.5)
    SELECTED_ZONE_RELATIVE_AIR_VELOCITY=st.slider("SELECTED_ZONE_RELATIVE_AIR_VELOCITY", 0.1, 1.0,0.25,step=0.05)
    METABOLIC_EQUIVALENT_MER=st.slider("METABOLIC_EQUIVALENT_MER", 0.7, 4.0,1.0,step=0.1)
    CLOTHING_INSULATION_CLO=st.slider("METABOLIC_EQUIVALENT_MER", 0.36, 1.0,0.61,step=0.05)
    CURRENT_LOCATION=get_location()
    SITE_LOCATION_OF_THE_UNIT=st.selectbox(" SITE_LOCATION_OF_THE_UNIT",[CURRENT_LOCATION,'bangalore','mumbai', 'chennai'],)

    
    AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW=st.number_input("AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW")
    SELECTED_MOTOR_EFFICIENCY=st.slider("SELECTED_MOTOR_EFFICIENCY", 50.0, 100.0,90.0,step=0.5)
    SELECTED_MOTOR_POWER_FACTOR=st.slider("SELECTED_MOTOR_POWER_FACTOR", 0.7, 1.0,0.89,step=0.05)
    SELECTED_RATED_CURRENT=st.number_input("SELECTED_RATED_CURRENT")
    SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS",value=16)
    SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIU",value=7)
    SELECTED_CHILLED_WATER_FLOWRATE_USGPM=st.number_input("SELECTED_CHILLED_WATER_FLOWRATE_USGPM",value=50)
    
    w=[DATE_OF_INSPECTION,AHU_TAG_NAME,
       AHU_DESIGN_CFM,AHU_DESIGN_ESP_IN_MMWG,SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS,SELECTED_INDOOR_RELATIVE_HUMIDITY,SELECTED_ZONE_RELATIVE_AIR_VELOCITY,METABOLIC_EQUIVALENT_MER,CLOTHING_INSULATION_CLO,SITE_LOCATION_OF_THE_UNIT,AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW,SELECTED_MOTOR_EFFICIENCY,SELECTED_MOTOR_POWER_FACTOR,SELECTED_RATED_CURRENT,SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS, SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS,SELECTED_CHILLED_WATER_FLOWRATE_USGPM]
    
    
    st.header("SITE AIR VELOCITY FPM READINGS")
    acfm=st.number_input("ENTER NO OF VELOCITY READING GOING TO TAKEN AT SITE>>>", min_value=1, step=1, value=5)
    cfmreading=[]
    for i in range(1,acfm+1):
        user_input = st.number_input(f"ENTER {i} READING OUT OF {acfm} READINGS", min_value=0, step=1)
        cfmreading.append(user_input)
    
    st.success(f"THE ENTERED FPM VALUE OF THE UNITS ARE>>> {cfmreading} .")
    avgfpm=sum(cfmreading)/len(cfmreading)
    st.success(f"THE AVERAGE FPM VALUE OF THE READINGS ARE>>>{avgfpm}")
    
    
    # OpenWeatherMap API endpoint and API key
    weather_api_endpoint = "http://api.openweathermap.org/data/2.5/weather"
    api_key = "8af1a752ec6e57631f899729ded7eba5"  # Replace this with your actual OpenWeatherMap API key

    params = {'q':  get_location(),
            'appid': api_key,
            'units': 'metric'  }
    response = requests.get(weather_api_endpoint, params=params)
    if response.status_code == 200:
        weather_data = response.json()
        main = weather_data['main']
        weather = weather_data['weather'][0]['description']
        temperature = main['temp']
        atm= main['pressure']
        humidity = main['humidity']
        lon=(weather_data['coord']['lon'])
        lat=(weather_data['coord']['lat'])
        win=(weather_data['wind']['speed'])
        city=(weather_data['name'])
    else:
        print("Error fetching weather data.")

    ENTER_FILTER_FACE_AREA_WIDTH_IN_mm=st.number_input("ENTER_FILTER_FACE_AREA_WIDTH_IN_mm")
    ENTER_FILTER_FACE_AREA_HEIGHT_IN_mm=st.number_input("ENTER_FILTER_FACE_AREA_HEIGHT_IN_mm")
    filterwidth=ENTER_FILTER_FACE_AREA_WIDTH_IN_mm
    filterheight=ENTER_FILTER_FACE_AREA_HEIGHT_IN_mm

    sqft=filterwidth/1000*filterheight/1000*10.764
    st.success(f"AHU'S MEASURED FILTER FACE AREA IS {sqft} SQFT")
    
    st.header("TEMPERATURE/FLOWRATE/VOLTAGE/CURRENT MEASUREMENT.")
    
    RETURN_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("RETURN_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0,26.0,step=0.05)
    RETURN_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("RETURN_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0,24.0,step=0.05)
    SUPPLY_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("SUPPLY_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0,22.0,step=0.05)
    SUPPLY_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("SUPPLY_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0,20.0,step=0.05)
    MEASURED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("MEASURED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS",value=6)
    MEASURED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS=st.number_input("MEASURED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS",value=16)
    MEASURED_CHILLED_WATER_FLOWRATE_USGPM=st.number_input("MEASURED_CHILLED_WATER_FLOWRATE_USGPM",value=26)
    MEASURED_VOLTAGE_R_AND_Y=st.number_input("MEASURED_VOLTAGE_R_AND_Y",value=415)
    MEASURED_VOLTAGE_Y_AND_B=st.number_input("MEASURED_VOLTAGE_Y_AND_B",value=415)
    MEASURED_VOLTAGE_B_AND_R=st.number_input("MEASURED_VOLTAGE_B_AND_R",value=415)
    MEASURED_CURRENT_R=st.number_input("MEASURED_CURRENT_R",value=2)
    MEASURED_CURRENT_Y=st.number_input("MEASURED_CURRENT_Y",value=2)
    MEASURED_CURRENT_B=st.number_input("MEASURED_CURRENT_B",value=2)

    MEASURED_ZONE_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS=st.slider("MEASURED_ZONE_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS", 10.0, 60.0,24.0,step=0.05)
    MEASURED_ZONE_RELATIVE_HUMIDITY=st.slider("MEASURED_ZONE_RELATIVE_HUMIDITY", 10.0, 100.0,55.0,step=0.05)
    MEASURED_MEAN_RADIANT_TEMPERATURE=st.number_input("MEASURED_MEAN_RADIANT_TEMPERATURE",value=20)
    
    indbt=RETURN_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS
    inwbt=RETURN_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS
    outdbt=SUPPLY_AIR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS
    outwbt=SUPPLY_AIR_WET_BULB_TEMPERATURE__DEGREE_CELSIUS



    atmpressure=atm*100
    psychrolib.SetUnitSystem(psychrolib.SI)
    inair=psychrolib.CalcPsychrometricsFromTWetBulb(indbt, inwbt, atmpressure)
    outair=psychrolib.CalcPsychrometricsFromTWetBulb(outdbt, outwbt, atmpressure)

    table_header1= ["PSYCROMETRIC PARAMETER", "RETURN AIR", "SUPPLY AIR"]
    table_data1 = [["Humidity ratio in kg_H₂O kg_Air",inair[0],outair[0]],
    ["Dew-point temperature in  °C [SI]",inair[1],outair[1]],
    ["Relative humidity in [0,1] ",inair[2],outair[2]],
    ["Partial pressure of water vapor in moist air in  Pa [SI] ",inair[3],outair[3]],
    ["Moist air enthalpy in J kg⁻¹ [SI] ",inair[4],outair[4]],
    ["Specific volume of moist air in  in m³ kg⁻¹ [SI]" ,inair[5],outair[5]],
    ["Degree of saturation [unitless] ",inair[6],outair[6]]]


    data= {
    'PSYCROMETRIC PARAMETER':["DRY_BULB_TEMPERATURE__DEGREE_CELSIUS",
                            "WET_BULB_TEMPERATURE__DEGREE_CELSIUS",
                            "Humidity ratio in kg_H₂O kg_Air",
                            "Dew-point temperature in  °C [SI]",
                            "Relative humidity in [0,100]",
                            "Partial pressure of water vapor in moist air in  Pa [SI]",
                            "Moist air enthalpy in J kg⁻¹ [SI]",
                            "Specific volume of moist air in  in m³ kg⁻¹ [SI]",
                            "Degree of saturation [unitless]"],
    'RETURN AIR':[indbt,
                inwbt,
                inair[0],
                inair[1],
                inair[2]*100,
                inair[3],
                inair[4],
                inair[5],
                inair[6]],
    'SUPPLY AIR':[outdbt,
                outwbt,
                outair[0],
                outair[1],
                outair[2]*100,
                outair[3],
                outair[4],
                outair[5],
                outair[6]]}

    df=pd.DataFrame(data,index=None)


    table_header1=['PSYCROMETRIC PARAMETER','RETURN AIR','SUPPLY AIR']
    table_data1= [
            ["DRY_BULB_TEMPERATURE__DEGREE_CELSIUS" , indbt , outdbt ],
            ["WET_BULB_TEMPERATURE__DEGREE_CELSIUS",inwbt,outwbt],
            ["Humidity ratio in kg_H₂O kg_Air",inair[0],outair[0]],
            ["Dew-point temperature in  °C [SI]",inair[1],outair[1]],
            ["Relative humidity in [0,100]",inair[2],outair[2]],
            ["Partial pressure of water vapor in moist air in  Pa [SI]",inair[3],outair[3]],
            ["Moist air enthalpy in J kg⁻¹ [SI]",inair[4],outair[4]],
            ["Specific volume of moist air in  in m³ kg⁻¹ [SI]",inair[5],outair[5]],
            ["Humidity ratio in kg_H₂O kg_Air",inair[6],outair[6]],
        ]
    document = Document()

    document.add_heading('AHU FUNCTIONAL & PERFORMANCE TEST REPORT', 0)

    ##AHU SELECTION DATA

    w1=["DATE_OF_INSPECTION",
    "AHU_TAG_NAME",
    "AHU_DESIGN_CFM",
    "AHU_DESIGN_ESP_IN_MMWG",
    "SELECTED_INDOOR_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_INDOOR_RELATIVE_HUMIDITY",
    "SELECTED_ZONE_RELATIVE_AIR_VELOCITY",
    "METABOLIC_EQUIVALENT_MER",
    "CLOTHING_INSULATION_CLO",
    "SITE_LOCATION_OF_THE_UNIT",
    "AHU_SELECTED_OUTPUT_RATED_POWER_IN_KW",
    "SELECTED_MOTOR_EFFICIENCY",
    "SELECTED_MOTOR_POWER_FACTOR",
    "SELECTED_RATED_CURRENT",
    "SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS",
    "SELECTED_CHILLED_WATER_FLOWRATE_USGPM"]

    document.add_heading('AHU SELECTION DATA', level=1)
    for i in range(len(w)):
        document.add_paragraph(str(w1[i])+" =  " +str(w[i]), style='List Bullet')



    document.add_heading('AMBIENT CONDITION DATA', level=1)
    
    amb=[SITE_LOCATION_OF_THE_UNIT.upper(),lon,lat,temperature,humidity,weather,atm,win]
    ambtxt=["SITE_LOCATION_OF_THE_UNIT","LONGITUDE","LATITUDE","TEMPERATURE","HUMIDITY","WEATHER CONDITION","ATMOSPHERIC PRESSURE","AIR SPEED"]
    ambtxtunit=["","degrees","degrees","°C,","%"," ","hpa","m/s"]


    for i in range(len(amb)):
        document.add_paragraph(ambtxt[i]+" =  "  +str(amb[i])+"  " +str(ambtxtunit[i]), style='List Bullet')

    ## BLANK LINES
    for i in range(4):
        document.add_paragraph()
    
    q1=["DRY_BULB_TEMPERATURE__DEGREE_CELSIUS",
                            "WET_BULB_TEMPERATURE__DEGREE_CELSIUS",
                            "Humidity ratio in kg_H₂O kg_Air",
                            "Dew-point temperature in  °C [SI]",
                            "Relative humidity in [0,100]",
                            "Partial pressure of water vapor in moist air in  Pa [SI]",
                            "Moist air enthalpy in J kg⁻¹ [SI]",
                            "Specific volume of moist air in  in m³ kg⁻¹ [SI]",
                            "Degree of saturation [unitless]"]
    RETURN_AIR=[indbt,
                inwbt,
                inair[0],
                inair[1],
                inair[2]*100,
                inair[3],
                inair[4],
                inair[5],
                inair[6]]
    SUPPLY_AIR=[outdbt,
                outwbt,
                outair[0],
                outair[1],
                outair[2]*100,
                outair[3],
                outair[4],
                outair[5],
                outair[6]]


    document.add_heading('MEASURED SUPPLY AIR PSYCHROMETRIC PARAMETERS', level=1)

    for i in range(9):
        document.add_paragraph(q1[i]+" =  "  +str(SUPPLY_AIR[i]), style='List Bullet')

    document.add_heading('MEASURED RETURN AIR PSYCHROMETRIC PARAMETERS', level=1)

    for i in range(9):
        document.add_paragraph(q1[i]+" =  "  +str(RETURN_AIR[i]), style='List Bullet')    
    
    ##PSYCHRO CHART

    document.add_heading('PSYCHROMETRIC PLOTTING OF AIR PARAMETER AT SEA LEVEL OF  '+ SITE_LOCATION_OF_THE_UNIT.upper(), level=1)


    # plot a simple psychrometric chart


    psychrolib.SetUnitSystem(psychrolib.SI)

    pressure = atmpressure

    t_array = np.arange(5, 45, 0.1)
    rh_array = np.arange(0, 1.1, 0.1)
    enthalpy_array = np.arange(0, 120000, 10000)
    hr_hor_lines = np.arange(0.005, 0.03, 0.005)
    twb_array = np.arange(-10, 45, 5)

    f, ax = plt.subplots()

    # plot constant relative humidity lines
    for rh in rh_array:
        hr_array = []
        for t in t_array:
            hr = psychrolib.GetHumRatioFromRelHum(t, rh, pressure)
            hr_array.append(hr)
        ax.plot(t_array, hr_array, 'k')

    for twb in twb_array:
        hr_array = []
        t_plot_array = []
        for t in t_array:
            if twb <= t:
                # print(twb, t)
                hr = psychrolib.GetHumRatioFromTWetBulb(t, twb, pressure)
                hr_array.append(hr)
                t_plot_array.append(t)
        ax.plot(t_plot_array, hr_array, 'b')

    # Define the components of the vector

    x1_component = outdbt
    y1_component = psychrolib.GetHumRatioFromTWetBulb(outdbt, outwbt, pressure)


    x_component = indbt
    y_component = psychrolib.GetHumRatioFromTWetBulb(indbt, inwbt, pressure)

    # Plot the vector using plt.plot()
    ax.plot([x1_component, x_component], [y1_component, y_component], color='green', linewidth=2, marker='o', markersize=8)



    # Add text with an arrow pointing to the target point
    ax.annotate('Supply Air point', xy=(x1_component-2, y1_component-0.002))
    ax.annotate('Return Air point', xy=(x_component+2, y_component+0.002))

    ax.set(ylim=(0, 0.025), xlim=(10, 40), ylabel=r"Humidity Ratio [$kg_{water}/kg_{dry air}$]", xlabel="Dry-bulb Temperature [°C]")
    ax.yaxis.tick_right()
    ax.yaxis.set_label_position("right")
    plt.title('Plotting onsite parameters in Pychrometric chart')
    plt.tight_layout()

    plt.savefig('table.png')

    document.add_picture('table.png', width=Inches(1.25*5))    
    
#####CHILLED WATER PARAMETERS

    document.add_heading('CHILLED WATER PARAMETERS', level=1)

    document.add_paragraph(str("MEASURED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS")+
                        " = > " +
                        str(MEASURED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS), style='List Bullet')

    document.add_paragraph(str("MEASURED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS")+
            " = >  "  +
            str(MEASURED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS), style='List Bullet')

    document.add_paragraph(str("MEASURED_CHILLED_WATER_FLOWRATE_USGPM")+
            " = >  "  +
            str(MEASURED_CHILLED_WATER_FLOWRATE_USGPM), style='List Bullet')


    #####ELECTRICAL PARAMETERS

    document.add_heading('ELECTRICAL PARAMETERS', level=1)

    document.add_paragraph(str("MEASURED_VOLTAGE_R_AND_Y")+
            " = >  "  +
            str(MEASURED_VOLTAGE_R_AND_Y), style='List Bullet')

    document.add_paragraph(str("MEASURED_VOLTAGE_Y_AND_B")+
            " = >  "  +
            str(MEASURED_VOLTAGE_Y_AND_B), style='List Bullet')

    document.add_paragraph(str("MEASURED_VOLTAGE_B_AND_R")+
            " = >  "  +
            str(MEASURED_VOLTAGE_B_AND_R), style='List Bullet')

    document.add_paragraph(str("MEASURED_CURRENT_R")+
            " = >  "  +
            str(MEASURED_CURRENT_R), style='List Bullet')

    document.add_paragraph(str("MEASURED_CURRENT_Y")+
            " = >  "  +
            str(MEASURED_CURRENT_Y), style='List Bullet')

    document.add_paragraph(str("MEASURED_CURRENT_B")+
            " = >  "  +
            str(MEASURED_CURRENT_B), style='List Bullet')

    v=[MEASURED_VOLTAGE_R_AND_Y,MEASURED_VOLTAGE_Y_AND_B,MEASURED_VOLTAGE_B_AND_R]
    i=[MEASURED_CURRENT_R,MEASURED_CURRENT_Y,MEASURED_CURRENT_B]

    ikw=(1.732*(sum(v)/3)*(sum(i)/3)*SELECTED_MOTOR_POWER_FACTOR)/1000

    document.add_paragraph(str("INPUT POWER CONSUMPTION")+
            " = >  "  +
            str(ikw), style='List Bullet')

    document.add_paragraph(str("SELECTED_MOTOR_EFFICIENCY")+
            " = >  "  +
            str(SELECTED_MOTOR_EFFICIENCY), style='List Bullet')

    document.add_paragraph(str("MOTOR_OUTPUT SHAFT POWER")+
            " = >  "  +
            str(ikw*SELECTED_MOTOR_EFFICIENCY/100), style='List Bullet')


    ##### ZONE DATA

    document.add_heading('ZONE AIR PARAMETERS', level=1)

    document.add_paragraph(str("MEASURED_ZONE_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS")+
            " = >  "  +
            str(MEASURED_ZONE_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS), style='List Bullet')

    document.add_paragraph(str("MEASURED_ZONE_RELATIVE_HUMIDITY")+
            " = >  "  +
            str(MEASURED_ZONE_RELATIVE_HUMIDITY), style='List Bullet')

    document.add_paragraph(str("MEASURED_MEAN_RADIANT_TEMPERATURE")+
            " = >  "  +
            str(MEASURED_MEAN_RADIANT_TEMPERATURE), style='List Bullet')

    document.add_paragraph(str("THE ENTERED FPM VALUE OF THE UNITS ARE>>>")+
            " = >  "  +
            str(cfmreading), style='List Bullet')

    document.add_paragraph(str("THE AVERAGE FPM VALUE OF THE READINGS ARE>>>")+
            " = >  "  +
            str(avgfpm), style='List Bullet')

    document.add_paragraph(str("ENTER_FILTER_FACE_AREA_WIDTH_IN_mm>>>")+
            " = >  "  +
            str(ENTER_FILTER_FACE_AREA_WIDTH_IN_mm), style='List Bullet')

    document.add_paragraph(str("ENTER_FILTER_FACE_AREA_HEIGHT_IN_mm>>>")+
            " = >  "  +
            str(ENTER_FILTER_FACE_AREA_HEIGHT_IN_mm), style='List Bullet')

    document.add_paragraph(str("AHU'S MEASURED FILTER FACE AREA IN SQFT>>>")+
            " = >  "  +
            str(sqft), style='List Bullet')



    inh=inair[4]/1000*0.429923
    outh=outair[4]/1000*0.429923
    dh=inh-outh
    # print(inh,outh)
    cfm=sqft*avgfpm
    # print(f"AHU'S MEASURED AIR FLOW RATE IS {cfm} CFM")
    tr=4.5*cfm*dh/12000
    # print(dh)

    document.add_heading('AHU CALCULATION OUTPUTS', level=1)


    document.add_paragraph(str("AHU'S MEASURED AIR FLOW RATE IN CFM>>>")+
            " = >  "  +
            str(cfm), style='List Bullet')


    document.add_paragraph(str("AHU'S MEASURED TON OF REFEGERANTION>>>")+
            " = >  "  +
            str(tr), style='List Bullet')


    zonedbt = MEASURED_ZONE_DRY_BULB_TEMPERATURE__DEGREE_CELSIUS
    zonemrt = MEASURED_MEAN_RADIANT_TEMPERATURE
    zonerh = MEASURED_ZONE_RELATIVE_HUMIDITY
    v = SELECTED_ZONE_RELATIVE_AIR_VELOCITY
    met =METABOLIC_EQUIVALENT_MER
    clo =CLOTHING_INSULATION_CLO
    # calculate relative air speed
    v_r = v_relative(v=v, met=met)
    # calculate dynamic clothing
    clo_d = clo_dynamic(clo=clo, met=met)
    results = pmv_ppd(tdb=zonedbt, tr=zonemrt, vr=v_r, rh=zonerh, met=met, clo=clo_d)
    # print("Predicted Percentage of Dissatisfied (PPD) is "+str(results["ppd"]))
    # print("Predicted Mean Vote (PMV) is "+str(results["pmv"]))


    document.add_heading('ASHARE 55=>>THERMAL COMFORT CALCULATION', level=1)


    document.add_paragraph(str("PREDICTED PERCENTAGE OF DISSATISFIED (PPD)=")+
            " = >  "  +
            str(results["ppd"]), style='List Bullet')


    document.add_paragraph(str("PREDICTED MEAN VOTE (PMV) ==")+
            " = >  "  +
            str(results["pmv"]), style='List Bullet')


    ##PERFROMACE devSTAS

    chwin=((MEASURED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS-SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS)/SELECTED_CHILLED_WATER_INLET_TEMPERATURE__DEGREE_CELSIUS)*100
    chwout=((MEASURED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS-SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS)/SELECTED_CHILLED_WATER_OUTLET_TEMPERATURE__DEGREE_CELSIUS)*100
    # ahucfm=(abs(AHU_DESIGN_CFM-cfm)/AHU_DESIGN_CFM)
    # curr=(abs((sum(i)/3)-SELECTED_RATED_CURRENT )/SELECTED_RATED_CURRENT)

    pf=[chwin,chwout]
    pft=["CHILLED WATER INLET DEVIATION FROM SELCTION",
        "CHILLED WATER OUTLET DEVIATION FROM SELCTION"]
        # "AHU AIR FLOW DEVIATION FROM DESIGN CFM",
        # "CURRNT DEVIATION FROM RATED"]

    document.add_heading('AHU PERFORMANCE DEVIATION STATS', level=1)

    for i in range(len(pf)):
        document.add_paragraph(pft[i]+" =  " +str(pf[i])+"  %", style='List Bullet')   
        
    ##SIGNOFF

    document.add_heading('SIGNOFF......', level=1)

    records = (
        ('NAME:   ', 'TEAM:   ', 'SIGNATURE:   '),
        ('NAME:   ', 'TEAM:   ', 'SIGNATURE:   '),
        ('NAME:   ', 'TEAM:   ', 'SIGNATURE:   '),
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ATTENDEES'
    hdr_cells[1].text = 'PARTIES'
    hdr_cells[2].text = 'SIGNATURE'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = qty
        row_cells[1].text = id
        row_cells[2].text = desc
    
    if st.button("Generate AHU functional Report"):
        # Save the pdf with name
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        st.success("Report generated successfully!")
        st.download_button(label="Download AHU CX Report",
        data=buffer, file_name=AHU_TAG_NAME+'_CX_REPORT.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def main():
    st.sidebar.title("BIJU_AI")
    st.sidebar.image("https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcR-zAHtwRkbhXFAETMN3tqgGWSuaEbYcGOSyw&usqp=CAU")
    # Create a navigation menu in the sidebar
    app_mode = st.sidebar.radio("Go to", ["HOME", "AHU PRE COMMISSIONING REPORT", "SNAG_LIST_LOGGER_AI","AHU PERFORMANCE REPORT"])
    
    # Display the selected page
    if app_mode == "HOME":
        home()
    elif app_mode == "AHU PRE COMMISSIONING REPORT":
        form()
    elif app_mode == "SNAG_LIST_LOGGER_AI":
        data_display()
    elif app_mode == "AHU PERFORMANCE REPORT":
        loop()

if __name__ == "__main__":
    main()
